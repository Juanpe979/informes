#!/usr/bin/env python3
"""Informe calidad semana 2-8 marzo 2026 — estilo Notion"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

TEAL   = 'E8F1EC'
ORANGE = 'FBE9D4'
RED    = 'FCE8E7'
BLUE   = 'E3F1FB'

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D0CFC9')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x37, 0x35, 0x2F)
        run.font.size = Pt(18)
        run.font.bold = True

def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x37, 0x35, 0x2F)
        run.font.size = Pt(13)
        run.font.bold = True

def body(text, italic=False, parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if parts:
        for txt, bold in parts:
            run = p.add_run(txt)
            run.bold = bold
            run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.italic = italic
        run.font.size = Pt(11)

def bullet(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'D0CFC9')
    pBdr.append(bottom)
    pPr.append(pBdr)

def callout(parts, bg='FBE9D4', border='D27B2D'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:color'), border)
    pBdr.append(left)
    pPr.append(pBdr)
    for txt, bold in parts:
        run = p.add_run(txt)
        run.bold = bold
        run.font.size = Pt(11)

def make_table(headers, rows, row_colors):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_borders(t)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(cell, BLUE)
    for r_idx, (row_data, bg) in enumerate(zip(rows, row_colors)):
        row = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(cell, bg)
    doc.add_paragraph()

# ── DOCUMENTO ─────────────────────────────────────────────────────────────────

# Título
title = doc.add_heading('Review: Respuestas Automáticas Adri', 0)
for run in title.runs:
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x37, 0x35, 0x2F)

subtitle = doc.add_paragraph()
run = subtitle.add_run('Semana 2 – 8 de marzo de 2026')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x7D, 0x7A, 0x75)
subtitle.paragraph_format.space_after = Pt(10)

callout([
    ('Review de las ', False),
    ('18 conversaciones marcadas como "Mal"', True),
    (' durante la semana del 2 al 8 de marzo. Se identifican ', False),
    ('8 patrones de fallo', True),
    (', de los cuales 2 son urgentes por contener información incorrecta '
     'entregada al usuario. Los fallos de triaje se excluyen de este análisis.', False),
])

doc.add_paragraph()
divider()

# ── 1. RESUMEN ────────────────────────────────────────────────────────────────
h1('Resumen')
body('', parts=[
    ('18 casos ', True),
    ('de respuesta incorrecta esta semana. Cuatro patrones principales: '
     '(1) ', False),
    ('información incorrecta o desactualizada en prompts', True),
    (' — guardería, eficiencia energética, donativos, indemnizaciones (6 casos); '
     '(2) ', False),
    ('respuestas genéricas por falta de variables de contexto', True),
    (' — CCAA y yennefer (2 casos); (3) ', False),
    ('errores de año fiscal', True),
    (' en las respuestas (3 casos); (4) ', False),
    ('fallos de comportamiento del agente', True),
    (' — alucinaciones, preguntas innecesarias, scope no aclarado (4 casos). '
     'Al menos 6 de los 18 son corregibles con cambios directos en el prompt.', False),
])
divider()

# ── 2. FALLOS DETECTADOS ──────────────────────────────────────────────────────
h1('Fallos detectados')

make_table(
    headers=['Fallo', 'Nº casos', 'Prompt afectado', 'Impacto'],
    rows=[
        ['Información incorrecta en prompt: guardería (centro autorizado)', '2', 'Deducciones', 'Alto — respuesta factualmente errónea'],
        ['Información desactualizada: eficiencia energética', '2', 'Deducciones', 'Alto — respuesta incorrecta'],
        ['Error de año fiscal en la respuesta', '3', 'Múltiples', 'Alto — genera desconfianza'],
        ['Reglas incompletas: mutualidades, dos pagadores, indemnizaciones', '3', 'Trabajo', 'Medio — info incompleta'],
        ['Comportamiento: alucinaciones, seguimiento innecesario, scope sin aclarar', '4', 'System prompt / Experto otros ingresos', 'Medio-alto'],
        ['Respuesta genérica por falta de variable de contexto (CCAA / yennefer)', '2', 'Deducciones autonómicas', 'Medio'],
        ['Donativos desactualizados', '1', 'Deducciones', 'Alto — info incorrecta'],
        ['Clasificación multietiqueta deficiente', '1', 'Clasificador', 'Medio'],
        ['Deducción gimnasio Andalucía (prompt correcto, sin desplegar)', '1', 'Infraestructura', 'Bajo — no es fallo de prompt'],
    ],
    row_colors=[RED, RED, RED, ORANGE, ORANGE, ORANGE, RED, ORANGE, TEAL]
)

divider()

# ── 3. PROBLEMAS ESPECÍFICOS ──────────────────────────────────────────────────
h1('Problemas específicos')

h2('1. Información incorrecta sobre guardería (URGENTE)')
body('Registros: 9 (06/03), 13 (08/03)')
body('La IA indica que la guardería debe ser un "centro autorizado" para aplicar la deducción. '
     'Esto es incorrecto: la deducción aplica durante todo el año escolar y no exige autorización del centro. '
     'El error tiene origen directo en el prompt. Al repetirse en dos días distintos, '
     'confirma que es un fallo sistemático.')

h2('2. Eficiencia energética: deducciones inexistentes en el prompt')
body('Registros: 1 (04/03), 3 (01/03)')
body('El prompt incluye las dos primeras modalidades de deducción por eficiencia energética '
     '(vivienda individual) que ya no están vigentes. Solo debe quedar la tercera modalidad '
     '(edificios completos). El evaluador lo confirma explícitamente.')

h2('3. Error de año fiscal')
body('Registros: 11 (02/03), 17 (02/03), 20 (05/03)')
body('Tres casos donde la IA referencia el año equivocado. En el registro 17 indica que '
     '"la declaración ya está presentada" y redirige al usuario a 2025, cuando debería '
     'hablar del ejercicio en curso. En el 20, la teoría es correcta pero el año mencionado es 2024. '
     'Los prompts no tienen anclado de forma explícita el ejercicio fiscal 2025.')

h2('4. Reglas fiscales incompletas o mal explicadas')
body('Registros: 6 (02/03), 19 (06/03), 18 (05/03)')
bullet('Mutualidades (reg. 6): el prompt de trabajo no incluye que las cantidades recibidas de mutualidades pueden estar exentas.')
bullet('Regla de dos pagadores (reg. 19): la explicación actual no es suficientemente clara para el usuario.')
bullet('Indemnizaciones (reg. 18): la IA no distingue correctamente entre indemnizaciones exentas y no exentas.')

h2('5. Comportamiento del agente')
body('Registros: 8 (04/03), 7 (02/03), 16 (02/03), 10 (04/03)')
bullet('Alucinación (reg. 8): la IA mencionó un "reporte de prensa" que no existe en el prompt ni en el sistema.')
bullet('Seguimiento innecesario (reg. 7): tras indicar que Taxdown no gestiona un trámite, continúa haciendo preguntas sobre él.')
bullet('Scope sin aclarar (reg. 16): cuando el usuario pregunta por un impuesto que no gestionamos, la IA no lo indica al inicio.')
bullet('Respuesta parcial (reg. 10): explica cómo tributaría una ayuda pero no aclara que Taxdown no la gestiona.')

h2('6. Variable de contexto no disponible')
body('Registros: 2 (04/03), 15 (05/03)')
body('En ambos casos el evaluador confirma que si el sistema tuviera la variable CCAA o yennefer '
     'la respuesta habría sido correcta. No es un fallo de conocimiento del modelo sino de '
     'arquitectura de contexto.')

h2('7. Donativos desactualizados')
body('Registro: 12 (04/03)')
body('El funcionamiento de los donativos en el prompt no está actualizado. Requiere revisión '
     'con el equipo fiscal antes de modificar.')

h2('8. Deducción gimnasio Andalucía — fallo de despliegue, no de prompt')
body('Registro: 14 (08/03)')
body('La deducción SÍ está en el prompt actualizado, pero ese prompt no está en producción. '
     'La IA responde con la versión anterior. Acción: despliegue del prompt, no cambio de contenido.')

divider()

# ── 4. CAMBIOS RECOMENDADOS ───────────────────────────────────────────────────
h1('Cambios recomendados')

h2('Prompt Deducciones — Corregir: guardería (URGENTE)')
body('Eliminar cualquier mención al requisito de "centro autorizado". Texto a añadir:')
callout([
    ('La deducción por guardería aplica durante todo el año escolar. No es necesario que el '
     'centro sea un centro autorizado. El requisito es que el menor esté matriculado y que '
     'los gastos estén justificados documentalmente.', False)
], bg='F0F0F0', border='AAAAAA')
body('Ejemplo de respuesta esperada:')
body('"La deducción por guardería no exige que el centro esté autorizado. Aplica durante todo '
     'el año escolar, siempre que puedas acreditar los gastos con factura o recibo del centro."', italic=True)

h2('Prompt Deducciones — Corregir: eficiencia energética (URGENTE)')
body('Eliminar las dos primeras modalidades (vivienda individual). Conservar solo la tercera (edificios).')
body('Texto a añadir como aclaración:')
callout([
    ('Las deducciones por obras de eficiencia energética en vivienda habitual ya no están vigentes '
     'en el ejercicio 2025. Solo aplica la deducción para obras de mejora en edificios completos.', False)
], bg='F0F0F0', border='AAAAAA')
body('Ejemplo de respuesta esperada:')
body('"En cuanto a eficiencia energética, actualmente solo está vigente la deducción para obras '
     'de mejora en edificios. Las deducciones por obras en vivienda individual ya no están en vigor."', italic=True)

h2('System prompt — Añadir: año del ejercicio fiscal (URGENTE)')
body('Añadir al inicio del system prompt o instrucciones generales:')
callout([
    ('El ejercicio fiscal al que se refieren todas las consultas es el ejercicio 2025, '
     'que se declara en la campaña de renta de 2026. Cuando el usuario no especifique el año, '
     'asumir siempre que pregunta por el ejercicio 2025. No referenciar ejercicios anteriores '
     'salvo que el usuario lo indique explícitamente.', False)
], bg='F0F0F0', border='AAAAAA')
body('Ejemplo de respuesta esperada:')
body('"Para la declaración del ejercicio 2025, que presentas en esta campaña de 2026, '
     'el límite es el siguiente..."', italic=True)

h2('Prompt Trabajo — Añadir: mutualidades exentas')
callout([
    ('Las prestaciones recibidas de mutualidades de previsión social pueden estar total o '
     'parcialmente exentas dependiendo de las aportaciones realizadas antes del 1 de enero de 1999. '
     'Indicarlo siempre que el usuario mencione ingresos de una mutualidad.', False)
], bg='F0F0F0', border='AAAAAA')

h2('Prompt Trabajo — Mejorar: regla de dos pagadores')
callout([
    ('Con dos o más pagadores, el límite para estar obligado a declarar no es 22.000 € sino '
     '15.000 €, siempre que del segundo pagador se hayan cobrado más de 1.500 € en total '
     'durante el año.', False)
], bg='F0F0F0', border='AAAAAA')
body('Ejemplo de respuesta esperada:')
body('"Si has tenido dos pagadores, el límite para estar obligado a declarar es 15.000 €, '
     'no 22.000 €, siempre que del segundo hayas cobrado más de 1.500 € en el año."', italic=True)

h2('Prompt Deducciones — Revisar: donativos')
body('Actualizar el funcionamiento de los donativos con los porcentajes vigentes para el '
     'ejercicio 2025. Pendiente de validación fiscal antes de modificar.')

h2('Prompt Trabajo — Revisar: indemnizaciones exentas')
body('Añadir un listado claro de qué indemnizaciones están exentas y cuáles no, con sus '
     'condiciones. Pendiente de validación con equipo fiscal.')

h2('System prompt — Añadir: comportamiento en trámites no gestionados')
callout([
    ('Cuando el usuario pregunte por un trámite o impuesto que Taxdown no gestiona, '
     'indicarlo en el primer mensaje de forma clara. No hacer preguntas de seguimiento '
     'sobre ese trámite una vez comunicado. Cerrar esa línea y ofrecer ayuda en lo que sí se gestiona.', False)
], bg='F0F0F0', border='AAAAAA')
body('Ejemplo de respuesta esperada:')
body('"Ese trámite no lo gestionamos desde Taxdown. Si tienes dudas sobre tu declaración '
     'de la renta, estaré encantado de ayudarte."', italic=True)

h2('Infraestructura — Desplegar: prompt con deducción gimnasio Andalucía')
body('No requiere cambio de contenido. El prompt actualizado ya incluye la deducción. '
     'Acción: despliegue en producción.')

divider()

# ── 5. RESUMEN VISUAL ─────────────────────────────────────────────────────────
h1('Resumen visual')

make_table(
    headers=['Prioridad', 'Acción', 'Registros'],
    rows=[
        ['Urgente', 'Corregir prompt guardería (eliminar requisito centro autorizado)', '9, 13'],
        ['Urgente', 'Eliminar deducciones de eficiencia energética incorrectas', '1, 3'],
        ['Urgente', 'Anclar año 2025 en system prompt', '11, 17, 20'],
        ['Esta semana', 'Añadir mutualidades exentas al prompt de trabajo', '6'],
        ['Esta semana', 'Mejorar explicación regla dos pagadores', '19'],
        ['Esta semana', 'Añadir instrucción de comportamiento: scope y sin seguimiento', '7, 10, 16'],
        ['Esta semana', 'Desplegar prompt con deducción gimnasio Andalucía', '14'],
        ['Próximas semanas', 'Revisar y actualizar donativos (requiere validación fiscal)', '12'],
        ['Próximas semanas', 'Revisar indemnizaciones exentas con equipo fiscal', '18'],
        ['Próximas semanas', 'Mejorar gestión multietiqueta en clasificador', '5'],
        ['Próximas semanas', 'Incorporar variables CCAA y yennefer al contexto del sistema', '2, 15'],
    ],
    row_colors=[
        RED, RED, RED,
        ORANGE, ORANGE, ORANGE, ORANGE, ORANGE,
        TEAL, TEAL, TEAL, TEAL,
    ]
)

output = '/Users/juanpedrodurillomartin/Documents/prompts/prompt_acceso_auto/informe_semana_2_8_marzo.docx'
doc.save(output)
print(f'Guardado en: {output}')
